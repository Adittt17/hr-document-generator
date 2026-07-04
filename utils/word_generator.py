from docx import Document


def replace_placeholder(doc, placeholder, value):
    """
    Mengganti placeholder pada seluruh paragraf.
    """

    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:

            for run in paragraph.runs:
                run.text = run.text.replace(
                    placeholder,
                    str(value)
                )


def replace_table_placeholder(doc, placeholder, value):
    """
    Mengganti placeholder pada tabel.
    """

    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                for paragraph in cell.paragraphs:

                    for run in paragraph.runs:

                        run.text = run.text.replace(
                            placeholder,
                            str(value)
                        )


def generate_word(template_path, output_path, data):

    doc = Document(template_path)

    for key, value in data.items():

        placeholder = "{{" + key + "}}"

        replace_placeholder(
            doc,
            placeholder,
            value
        )

        replace_table_placeholder(
            doc,
            placeholder,
            value
        )

    doc.save(output_path)